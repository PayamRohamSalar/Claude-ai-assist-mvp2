# path: D:\OneDrive\AI-Project\Claude-ai-assist-mvp2\shared_utils\file_utils.py

"""
Legal Assistant AI - File Utilities
Provides utilities for file operations, document reading, and text processing
"""

import os
import hashlib
import mimetypes
import shutil
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Union, Any
from datetime import datetime
import json
import csv

# Document processing imports
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import chardet
    CHARDET_AVAILABLE = True
except ImportError:
    CHARDET_AVAILABLE = False

from .constants import (
    SUPPORTED_DOCUMENT_FORMATS, MAX_FILE_SIZE, Messages,
    RAW_DATA_DIR, BACKUP_DIR, persian_to_english_digits
)
from .logger import get_logger


class FileInfo:
    """Container for file information"""
    
    def __init__(self, file_path: Union[str, Path]):
        self.path = Path(file_path)
        self.name = self.path.name
        self.stem = self.path.stem
        self.suffix = self.path.suffix.lower()
        self.size = self.path.stat().st_size if self.path.exists() else 0
        self.modified_time = datetime.fromtimestamp(
            self.path.stat().st_mtime
        ) if self.path.exists() else None
        self.mime_type = mimetypes.guess_type(str(self.path))[0]
        self.encoding = None
        self.hash = None
        
    def calculate_hash(self, algorithm: str = 'md5') -> str:
        """Calculate file hash"""
        if not self.path.exists():
            return ""
        
        hash_func = hashlib.new(algorithm)
        with open(self.path, 'rb') as f:
            for chunk in iter(lambda: f.read(4096), b""):
                hash_func.update(chunk)
        
        self.hash = hash_func.hexdigest()
        return self.hash
    
    def detect_encoding(self) -> Optional[str]:
        """Detect file encoding"""
        if not CHARDET_AVAILABLE or not self.path.exists():
            return None
        
        try:
            with open(self.path, 'rb') as f:
                raw_data = f.read(10000)  # Read first 10KB
                result = chardet.detect(raw_data)
                self.encoding = result.get('encoding')
                return self.encoding
        except Exception:
            return None
    
    def to_dict(self) -> Dict[str, Any]:
        """Convert file info to dictionary"""
        return {
            'name': self.name,
            'stem': self.stem,
            'suffix': self.suffix,
            'size': self.size,
            'modified_time': self.modified_time.isoformat() if self.modified_time else None,
            'mime_type': self.mime_type,
            'encoding': self.encoding,
            'hash': self.hash,
            'path': str(self.path)
        }


class DocumentReader:
    """Document reader for various file formats"""
    
    def __init__(self):
        self.logger = get_logger("DocumentReader")
        self.supported_formats = {
            '.pdf': self._read_pdf,
            '.docx': self._read_docx,
            '.doc': self._read_doc,
            '.txt': self._read_text,
            '.json': self._read_json,
            '.csv': self._read_csv
        }
    
    def read_document(self, file_path: Union[str, Path]) -> Dict[str, Any]:
        """
        Read document content based on file extension
        
        Args:
            file_path: Path to the document
            
        Returns:
            Dictionary with content, metadata, and status
        """
        
        file_info = FileInfo(file_path)
        
        self.logger.info(
            f"Reading document: {file_info.name}",
            f"خواندن سند: {file_info.name}",
            file_size=file_info.size,
            file_type=file_info.suffix
        )
        
        # Validate file
        validation_result = self._validate_file(file_info)
        if not validation_result['valid']:
            return {
                'success': False,
                'content': '',
                'error': validation_result['error'],
                'file_info': file_info.to_dict()
            }
        
        # Read content based on file type
        reader_func = self.supported_formats.get(file_info.suffix)
        if not reader_func:
            error_msg = f"Unsupported file format: {file_info.suffix}"
            self.logger.error(error_msg, f"فرمت فایل پشتیبانی نمی‌شود: {file_info.suffix}")
            return {
                'success': False,
                'content': '',
                'error': error_msg,
                'file_info': file_info.to_dict()
            }
        
        try:
            content = reader_func(file_info)
            
            # Calculate file hash
            file_info.calculate_hash()
            
            self.logger.info(
                f"Document read successfully: {len(content)} characters",
                f"سند با موفقیت خوانده شد: {len(content)} کاراکتر"
            )
            
            return {
                'success': True,
                'content': content,
                'error': None,
                'file_info': file_info.to_dict(),
                'stats': {
                    'character_count': len(content),
                    'word_count': len(content.split()),
                    'line_count': content.count('\n') + 1
                }
            }
            
        except Exception as e:
            error_msg = f"Error reading document: {str(e)}"
            self.logger.error(error_msg, f"خطا در خواندن سند: {str(e)}")
            return {
                'success': False,
                'content': '',
                'error': error_msg,
                'file_info': file_info.to_dict()
            }
    
    def _validate_file(self, file_info: FileInfo) -> Dict[str, Any]:
        """Validate file before reading"""
        
        if not file_info.path.exists():
            return {
                'valid': False,
                'error': Messages.ERROR_FILE_NOT_FOUND
            }
        
        if file_info.size > MAX_FILE_SIZE:
            return {
                'valid': False,
                'error': f"File too large: {file_info.size} bytes (max: {MAX_FILE_SIZE})"
            }
        
        if file_info.suffix not in self.supported_formats:
            return {
                'valid': False,
                'error': Messages.ERROR_INVALID_FORMAT
            }
        
        return {'valid': True, 'error': None}
    
    def _read_pdf(self, file_info: FileInfo) -> str:
        """Read PDF file content"""
        
        if not PDF_AVAILABLE:
            raise ImportError("PyPDF2 not available for PDF reading")
        
        content = []
        
        with open(file_info.path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        content.append(f"--- صفحه {page_num + 1} ---\n{page_text}")
                except Exception as e:
                    self.logger.warning(
                        f"Error reading page {page_num + 1}: {str(e)}",
                        f"خطا در خواندن صفحه {page_num + 1}: {str(e)}"
                    )
        
        return '\n\n'.join(content)
    
    def _read_docx(self, file_info: FileInfo) -> str:
        """Read DOCX file content"""
        
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not available for DOCX reading")
        
        doc = DocxDocument(file_info.path)
        content = []
        
        # Read paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                content.append(paragraph.text)
        
        # Read tables if any
        for table in doc.tables:
            table_content = []
            for row in table.rows:
                row_content = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_content.append(cell.text.strip())
                if row_content:
                    table_content.append(' | '.join(row_content))
            
            if table_content:
                content.append('\n--- جدول ---\n' + '\n'.join(table_content))
        
        return '\n\n'.join(content)
    
    def _read_doc(self, file_info: FileInfo) -> str:
        """Read DOC file content (fallback to text reading)"""
        
        self.logger.warning(
            "DOC format has limited support, converting to text",
            "پشتیبانی محدود از فرمت DOC، تبدیل به متن"
        )
        
        # Try to read as text (this might not work well for binary DOC files)
        return self._read_text(file_info)
    
    def _read_text(self, file_info: FileInfo) -> str:
        """Read plain text file"""
        
        # Detect encoding
        encoding = file_info.detect_encoding() or 'utf-8'
        
        try:
            with open(file_info.path, 'r', encoding=encoding) as file:
                return file.read()
        except UnicodeDecodeError:
            # Fallback to utf-8 with error handling
            with open(file_info.path, 'r', encoding='utf-8', errors='ignore') as file:
                content = file.read()
                self.logger.warning(
                    "Used fallback encoding with error handling",
                    "از کدگذاری پشتیبان با مدیریت خطا استفاده شد"
                )
                return content
    
    def _read_json(self, file_info: FileInfo) -> str:
        """Read JSON file content"""
        
        with open(file_info.path, 'r', encoding='utf-8') as file:
            data = json.load(file)
            # Convert JSON to readable text
            return json.dumps(data, ensure_ascii=False, indent=2)
    
    def _read_csv(self, file_info: FileInfo) -> str:
        """Read CSV file content"""
        
        encoding = file_info.detect_encoding() or 'utf-8'
        content = []
        
        with open(file_info.path, 'r', encoding=encoding) as file:
            # Try to detect CSV dialect
            sample = file.read(1024)
            file.seek(0)
            
            try:
                dialect = csv.Sniffer().sniff(sample)
                reader = csv.reader(file, dialect)
            except csv.Error:
                reader = csv.reader(file)
            
            for row_num, row in enumerate(reader):
                if row and any(cell.strip() for cell in row):
                    row_text = ' | '.join(cell.strip() for cell in row if cell.strip())
                    content.append(f"ردیف {row_num + 1}: {row_text}")
        
        return '\n'.join(content)


class FileManager:
    """File management utilities"""
    
    def __init__(self):
        self.logger = get_logger("FileManager")
    
    def create_directory(self, directory: Union[str, Path], exist_ok: bool = True) -> bool:
        """Create directory with error handling"""
        
        try:
            Path(directory).mkdir(parents=True, exist_ok=exist_ok)
            self.logger.debug(f"Directory created: {directory}", f"دایرکتوری ایجاد شد: {directory}")
            return True
        except Exception as e:
            self.logger.error(f"Failed to create directory {directory}: {str(e)}")
            return False
    
    def copy_file(self, source: Union[str, Path], destination: Union[str, Path], 
                  overwrite: bool = False) -> bool:
        """Copy file with error handling"""
        
        source_path = Path(source)
        dest_path = Path(destination)
        
        if not source_path.exists():
            self.logger.error(f"Source file not found: {source_path}")
            return False
        
        if dest_path.exists() and not overwrite:
            self.logger.warning(f"Destination exists and overwrite is False: {dest_path}")
            return False
        
        try:
            # Create destination directory if needed
            dest_path.parent.mkdir(parents=True, exist_ok=True)
            shutil.copy2(source_path, dest_path)
            
            self.logger.info(
                f"File copied: {source_path} -> {dest_path}",
                f"فایل کپی شد: {source_path} -> {dest_path}"
            )
            return True
            
        except Exception as e:
            self.logger.error(f"Failed to copy file: {str(e)}")
            return False
    
    def move_file(self, source: Union[str, Path], destination: Union[str, Path]) -> bool:
        """Move file with error handling"""
        
        source_path = Path(source)
        dest_path = Path(destination)
        
        if not source_path.exists():
            self.logger.error(f"Source file not found: {source_path}")
            return False
        
        try:
            # Create destination directory if needed
            dest_path.parent.mkdir(parents=True, exist_ok=True)
            shutil.move(str(source_path), str(dest_path))
            
            self.logger.info(
                f"File moved: {source_path} -> {dest_path}",
                f"فایل منتقل شد: {source_path} -> {dest_path}"
            )
            return True
            
        except Exception as e:
            self.logger.error(f"Failed to move file: {str(e)}")
            return False
    
    def delete_file(self, file_path: Union[str, Path], backup: bool = True) -> bool:
        """Delete file with optional backup"""
        
        file_path = Path(file_path)
        
        if not file_path.exists():
            self.logger.warning(f"File not found for deletion: {file_path}")
            return True  # Consider it successful if file doesn't exist
        
        try:
            if backup:
                backup_path = BACKUP_DIR / "deleted_files" / f"{datetime.now().strftime('%Y%m%d_%H%M%S')}_{file_path.name}"
                if self.copy_file(file_path, backup_path):
                    self.logger.info(f"File backed up before deletion: {backup_path}")
            
            file_path.unlink()
            self.logger.info(f"File deleted: {file_path}", f"فایل حذف شد: {file_path}")
            return True
            
        except Exception as e:
            self.logger.error(f"Failed to delete file: {str(e)}")
            return False
    
    def get_directory_info(self, directory: Union[str, Path]) -> Dict[str, Any]:
        """Get directory information and statistics"""
        
        directory = Path(directory)
        
        if not directory.exists():
            return {
                'exists': False,
                'error': 'Directory not found'
            }
        
        try:
            files = []
            total_size = 0
            file_count = 0
            dir_count = 0
            
            for item in directory.rglob('*'):
                if item.is_file():
                    file_info = FileInfo(item)
                    files.append(file_info.to_dict())
                    total_size += file_info.size
                    file_count += 1
                elif item.is_dir():
                    dir_count += 1
            
            return {
                'exists': True,
                'path': str(directory),
                'file_count': file_count,
                'directory_count': dir_count,
                'total_size': total_size,
                'files': files,
                'created_time': datetime.fromtimestamp(directory.stat().st_ctime).isoformat(),
                'modified_time': datetime.fromtimestamp(directory.stat().st_mtime).isoformat()
            }
            
        except Exception as e:
            self.logger.error(f"Error getting directory info: {str(e)}")
            return {
                'exists': True,
                'error': str(e)
            }
    
    def find_files(self, directory: Union[str, Path], pattern: str = "*", 
                   extension: Optional[str] = None, recursive: bool = True) -> List[Path]:
        """Find files matching pattern"""
        
        directory = Path(directory)
        
        if not directory.exists():
            self.logger.error(f"Directory not found: {directory}")
            return []
        
        try:
            if recursive:
                search_func = directory.rglob
            else:
                search_func = directory.glob
            
            if extension:
                pattern = f"*.{extension.lstrip('.')}"
            
            files = [f for f in search_func(pattern) if f.is_file()]
            
            self.logger.debug(
                f"Found {len(files)} files matching pattern '{pattern}'",
                f"{len(files)} فایل با الگوی '{pattern}' یافت شد"
            )
            
            return files
            
        except Exception as e:
            self.logger.error(f"Error finding files: {str(e)}")
            return []
    
    def clean_filename(self, filename: str) -> str:
        """Clean filename for safe filesystem usage"""
        
        # Remove or replace problematic characters
        invalid_chars = ['<', '>', ':', '"', '/', '\\', '|', '?', '*']
        cleaned = filename
        
        for char in invalid_chars:
            cleaned = cleaned.replace(char, '_')
        
        # Handle Persian digits and spaces
        cleaned = persian_to_english_digits(cleaned)
        cleaned = ' '.join(cleaned.split())  # Normalize spaces
        
        # Limit length
        if len(cleaned) > 200:
            name, ext = os.path.splitext(cleaned)
            cleaned = name[:200-len(ext)] + ext
        
        return cleaned
    
    def get_file_type_stats(self, directory: Union[str, Path]) -> Dict[str, Dict[str, Any]]:
        """Get statistics by file type"""
        
        files = self.find_files(directory)
        stats = {}
        
        for file_path in files:
            file_info = FileInfo(file_path)
            ext = file_info.suffix or 'no_extension'
            
            if ext not in stats:
                stats[ext] = {
                    'count': 0,
                    'total_size': 0,
                    'files': []
                }
            
            stats[ext]['count'] += 1
            stats[ext]['total_size'] += file_info.size
            stats[ext]['files'].append(str(file_path))
        
        return stats


# Global instances
_document_reader: Optional[DocumentReader] = None
_file_manager: Optional[FileManager] = None

def get_document_reader() -> DocumentReader:
    """Get global document reader instance"""
    global _document_reader
    if _document_reader is None:
        _document_reader = DocumentReader()
    return _document_reader

def get_file_manager() -> FileManager:
    """Get global file manager instance"""
    global _file_manager
    if _file_manager is None:
        _file_manager = FileManager()
    return _file_manager

# Convenience functions
def read_document(file_path: Union[str, Path]) -> Dict[str, Any]:
    """Read document using global document reader"""
    return get_document_reader().read_document(file_path)

def create_directory(directory: Union[str, Path]) -> bool:
    """Create directory using global file manager"""
    return get_file_manager().create_directory(directory)

def copy_file(source: Union[str, Path], destination: Union[str, Path]) -> bool:
    """Copy file using global file manager"""
    return get_file_manager().copy_file(source, destination)

def get_file_info(file_path: Union[str, Path]) -> FileInfo:
    """Get file information"""
    return FileInfo(file_path)

def validate_file_type(file_path: Union[str, Path], allowed_extensions: List[str] = None) -> bool:
    """Validate if file type is allowed"""
    from .constants import validate_file_extension
    return validate_file_extension(str(file_path), allowed_extensions)